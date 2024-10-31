# backend/app/api/documents.py
import io
import time
from typing import List, Literal

from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, HTTPException, Response
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from sqlalchemy.orm import Session, joinedload
from sqlalchemy.sql import func

from ..database import get_db
from ..models.document import Document
from ..models.page import Page
from ..schemas.document import DocumentCreate, DocumentUpdate, Document as DocumentSchema, DocumentDetail
from ..services.cleanup import cleanup_service
from ..utils.logging import api_logger

router = APIRouter(prefix="/api/documents", tags=["documents"])


@router.get("/project/{project_id}", response_model=List[DocumentDetail])
async def list_project_documents(project_id: int, db: Session = Depends(get_db)):
    api_logger.info(f"Listing documents for project", extra={
        "project_id": project_id,
        "operation": "list_project_documents"
    })

    try:
        start_time = time.time()
        documents = db.query(Document).filter(Document.project_id == project_id).all()

        api_logger.debug("Retrieved documents from database", extra={
            "document_count": len(documents),
            "project_id": project_id
        })

        result = []
        for doc in documents:
            # Count pages for each document
            page_count = db.query(func.count(Page.id)).filter(Page.document_id == doc.id).scalar()
            doc_data = DocumentDetail.model_validate(doc)
            doc_data.page_count = page_count
            result.append(doc_data)

        execution_time = time.time() - start_time
        api_logger.info("Successfully listed project documents", extra={
            "project_id": project_id,
            "document_count": len(result),
            "execution_time_ms": round(execution_time * 1000, 2)
        })
        return result

    except Exception as e:
        api_logger.error("Error listing project documents", extra={
            "project_id": project_id,
            "error": str(e)
        })
        raise


@router.post("", response_model=DocumentSchema)
async def create_document(document: DocumentCreate, db: Session = Depends(get_db)):
    api_logger.info("Creating new document", extra={
        "project_id": document.project_id,
        "document_name": document.name
    })

    try:
        start_time = time.time()
        db_document = Document(**document.model_dump())
        db.add(db_document)
        db.commit()
        db.refresh(db_document)

        execution_time = time.time() - start_time
        api_logger.info("Successfully created document", extra={
            "document_id": db_document.id,
            "project_id": db_document.project_id,
            "execution_time_ms": round(execution_time * 1000, 2)
        })
        return db_document

    except Exception as e:
        api_logger.error("Error creating document", extra={
            "project_id": document.project_id,
            "document_name": document.name,
            "error": str(e)
        })
        db.rollback()
        raise


@router.get("/{document_id}", response_model=DocumentDetail)
async def get_document(document_id: int, db: Session = Depends(get_db)):
    api_logger.info("Retrieving document details", extra={
        "document_id": document_id
    })

    try:
        start_time = time.time()
        document = db.query(Document) \
            .options(joinedload(Document.pages)) \
            .filter(Document.id == document_id) \
            .first()

        if not document:
            api_logger.warning("Document not found", extra={
                "document_id": document_id
            })
            raise HTTPException(status_code=404, detail="Document not found")

        doc_data = DocumentDetail.model_validate(document)
        doc_data.page_count = len(document.pages)

        execution_time = time.time() - start_time
        api_logger.info("Successfully retrieved document", extra={
            "document_id": document_id,
            "page_count": doc_data.page_count,
            "execution_time_ms": round(execution_time * 1000, 2)
        })
        return doc_data

    except HTTPException:
        raise
    except Exception as e:
        api_logger.error("Error retrieving document", extra={
            "document_id": document_id,
            "error": str(e)
        })
        raise


@router.put("/{document_id}", response_model=DocumentSchema)
async def update_document(document_id: int, document: DocumentUpdate, db: Session = Depends(get_db)):
    api_logger.info("Updating document", extra={
        "document_id": document_id,
        "update_fields": list(document.model_dump(exclude_unset=True).keys())
    })

    try:
        start_time = time.time()
        db_document = db.query(Document).filter(Document.id == document_id).first()

        if not db_document:
            api_logger.warning("Document not found for update", extra={
                "document_id": document_id
            })
            raise HTTPException(status_code=404, detail="Document not found")

        # Log original values before update
        original_values = {
            field: getattr(db_document, field)
            for field in document.model_dump(exclude_unset=True).keys()
        }

        for field, value in document.model_dump(exclude_unset=True).items():
            setattr(db_document, field, value)

        db.commit()
        db.refresh(db_document)

        execution_time = time.time() - start_time
        api_logger.info("Successfully updated document", extra={
            "document_id": document_id,
            "original_values": original_values,
            "new_values": document.model_dump(exclude_unset=True),
            "execution_time_ms": round(execution_time * 1000, 2)
        })
        return db_document

    except HTTPException:
        raise
    except Exception as e:
        api_logger.error("Error updating document", extra={
            "document_id": document_id,
            "error": str(e)
        })
        db.rollback()
        raise


@router.delete("/{document_id}")
async def delete_document(document_id: int, db: Session = Depends(get_db)):
    api_logger.info("Deleting document", extra={"document_id": document_id})

    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            api_logger.warning("Document not found for deletion")
            raise HTTPException(status_code=404, detail="Document not found")

        # Delete document artifacts (includes pages)
        await cleanup_service.delete_document_artifacts(document, db)

        # Delete the document (will cascade delete pages in database)
        db.delete(document)
        db.commit()

        api_logger.info(f"Successfully deleted document {document_id}")
        return {"success": True}

    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        api_logger.error(f"Failed to delete document: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{document_id}/export")
async def export_document(
        document_id: int,
        format: Literal["pdf", "docx"] = "pdf",
        db: Session = Depends(get_db)
):
    api_logger.info("Starting document export", extra={
        "document_id": document_id,
        "format": format
    })

    try:
        # Verify document exists
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            api_logger.warning("Document not found", extra={"document_id": document_id})
            raise HTTPException(status_code=404, detail="Document not found")

        # Get all pages in order
        pages = db.query(Page).filter(
            Page.document_id == document_id
        ).order_by(Page.page_number).all()

        api_logger.debug("Retrieved pages", extra={
            "document_id": document_id,
            "page_count": len(pages)
        })

        if not pages:
            api_logger.warning("Document has no pages", extra={"document_id": document_id})
            raise HTTPException(status_code=400, detail="Document has no pages")

        # Verify pages have content
        pages_with_content = []
        for page in pages:
            text = page.formatted_text or page.extracted_text
            if text:
                pages_with_content.append({
                    'number': page.page_number,
                    'text': text
                })

        if not pages_with_content:
            api_logger.warning("No pages have text content", extra={"document_id": document_id})
            raise HTTPException(status_code=400, detail="No text content found in document")

        api_logger.debug("Processing pages with content", extra={
            "document_id": document_id,
            "pages_with_content": len(pages_with_content)
        })

        if format == "pdf":
            try:
                # Create PDF in memory
                buffer = io.BytesIO()
                doc = SimpleDocTemplate(
                    buffer,
                    pagesize=letter,
                    rightMargin=72,
                    leftMargin=72,
                    topMargin=72,
                    bottomMargin=72
                )

                styles = getSampleStyleSheet()
                content = []

                # Add document title
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=16,
                    spaceAfter=30
                )
                content.append(Paragraph(document.name, title_style))

                # Add page content
                for page_content in pages_with_content:
                    para = Paragraph(page_content['text'], styles['Normal'])
                    content.append(para)
                    content.append(Spacer(1, 12))

                # Build PDF
                doc.build(content)
                buffer.seek(0)

                api_logger.info("PDF export successful", extra={
                    "document_id": document_id,
                    "page_count": len(pages_with_content)
                })

                return Response(
                    content=buffer.getvalue(),
                    media_type="application/pdf",
                    headers={
                        "Content-Disposition": f'attachment; filename="{document.name}.pdf"'
                    }
                )

            except Exception as e:
                api_logger.error("PDF generation failed", extra={
                    "document_id": document_id,
                    "error": str(e)
                })
                raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

        elif format == "docx":
            try:
                # Create DOCX in memory
                doc = DocxDocument()
                doc.add_heading(document.name, 0)

                # Add page content
                for page_content in pages_with_content:
                    doc.add_paragraph(page_content['text'])
                    doc.add_paragraph()  # Add blank line between pages

                # Save to buffer
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                api_logger.info("DOCX export successful", extra={
                    "document_id": document_id,
                    "page_count": len(pages_with_content)
                })

                return Response(
                    content=buffer.getvalue(),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={
                        "Content-Disposition": f'attachment; filename="{document.name}.docx"'
                    }
                )

            except Exception as e:
                api_logger.error("DOCX generation failed", extra={
                    "document_id": document_id,
                    "error": str(e)
                })
                raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")

        else:
            api_logger.warning("Invalid format requested", extra={
                "document_id": document_id,
                "format": format
            })
            raise HTTPException(status_code=400, detail="Invalid format specified")

    except HTTPException:
        raise
    except Exception as e:
        api_logger.error("Unexpected error during export", extra={
            "document_id": document_id,
            "format": format,
            "error": str(e)
        })
        raise HTTPException(
            status_code=500,
            detail=f"Export failed: {str(e)}"
        )
